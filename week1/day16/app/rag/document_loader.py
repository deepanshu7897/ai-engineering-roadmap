from pathlib import Path

import fitz
from docx import Document as DocxDocument

from app.rag.document import Document
from app.rag.text_cleaner import TextCleaner


class DocumentLoader:

    def load(self, file_path: str) -> list[Document]:

        file_path = Path(file_path)

        suffix = file_path.suffix.lower()

        if suffix == ".pdf":
            return self._load_pdf(file_path)

        elif suffix == ".docx":
            return self._load_docx(file_path)

        elif suffix == ".txt":
            return self._load_txt(file_path)

        else:
            raise ValueError(
                f"Unsupported file type: {suffix}"
            )

    # -------------------------------------
    # PDF
    # -------------------------------------

    def _load_pdf(
        self,
        file_path: Path,
    ) -> list[Document]:

        pdf = fitz.open(file_path)

        documents = []

        for page_number, page in enumerate(pdf):

            text = page.get_text()

            text = TextCleaner.clean(text)

            if not text:
                continue

            documents.append(
                Document(
                    content=text,
                    metadata={
                        "source": file_path.name,
                        "page": page_number + 1,
                        "doc_type": "pdf",
                    },
                )
            )

        pdf.close()

        return documents

    # -------------------------------------
    # DOCX
    # -------------------------------------

    def _load_docx(
        self,
        file_path: Path,
    ) -> list[Document]:

        doc = DocxDocument(file_path)

        text = "\n".join(
            paragraph.text
            for paragraph in doc.paragraphs
        )

        text = TextCleaner.clean(text)

        return [
            Document(
                content=text,
                metadata={
                    "source": file_path.name,
                    "page": 1,
                    "doc_type": "docx",
                },
            )
        ]

    # -------------------------------------
    # TXT
    # -------------------------------------

    def _load_txt(
        self,
        file_path: Path,
    ) -> list[Document]:

        text = file_path.read_text(
            encoding="utf-8",
        )

        text = TextCleaner.clean(text)

        return [
            Document(
                content=text,
                metadata={
                    "source": file_path.name,
                    "page": 1,
                    "doc_type": "txt",
                },
            )
        ]